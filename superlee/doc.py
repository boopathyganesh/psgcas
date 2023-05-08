'''import docx
# Open the Word document
doc = docx.Document('../test.docx')

# Loop through all paragraphs in the document
for para in doc.paragraphs:
    # Check if the paragraph contains the '<name>' tag
    if '<name>' in para.text:
        # Replace the '<name>' tag with a specific name
        para.text = para.text.replace('<name>', 'Dhivya')

# Save the modified document
doc.save('modified_document.docx')'''
'''import docx
doc = docx.Document('../test.docx')

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if '<name>' in paragraph.text:
                   print('found date: ', paragraph.text)
                else:
                    print('NO')'''
'''import docx

# Open the Word document
doc = docx.Document('../test.docx')
print(doc.inline_shapes.has_text_frame)
# Loop through all the inline shapes in the document
for inline_shape in doc.inline_shapes:
    print(inline_shape)
    # Check if the inline shape is a text box
    if inline_shape.has_text_frame:
        # Loop through all paragraphs in the text box
        for para in inline_shape.text_frame.paragraphs:
            # Replace the string "old_text" with "new_text"
            para.text = para.text.replace("old_text", "new_text")

# Save the modified document
doc.save('modified_document.docx')'''


#mysql to xlsx
import pandas as pd
import mysql.connector
mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  password="",
  database="resumeai"
)

query = "SELECT * FROM `user`"

df = pd.read_sql(query, con=mydb)
print(df)
df.to_excel("output.xlsx", index=False)



'''from docx import Document

# Open the Word document
doc = Document('../test.docx')
count=0
# Loop through all elements in the document
for element in doc.element.body:
    count+=1
    print(count,element)
    # Check if the element is a textbox
    if element.tag.endswith('}txbxContent'):
        print('Textbox found')'''


'''import docx
from lxml import etree

# Open the Word document
doc = docx.Document('../test.docx')

# Extract the XML representation of the document
xml_bytes = etree.tostring(doc._element)

# Convert the XML to an lxml element tree
parser = etree.XMLParser(remove_blank_text=True)
xml_tree = etree.fromstring(xml_bytes, parser=parser)

# Define the namespace map for the XPath queries
ns_map = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'v': 'urn:schemas-microsoft-com:vml',
    'w10': 'urn:schemas-microsoft-com:office:word',
    'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing',
    'wp15': 'http://schemas.microsoft.com/office/word/2012/wordprocessingDrawing',
    'xdr': 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing',
    'xd': 'http://schemas.microsoft.com/office/2006/xdrawing',
    'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'o': 'urn:schemas-microsoft-com:office:office'
}

# Use XPath queries to extract the parts of the document you need
paragraphs = xml_tree.xpath('//w:p', namespaces=ns_map)
tables = xml_tree.xpath('//w:tbl', namespaces=ns_map)

#print(paragraphs)
'''
'''import zipfile
from lxml import etree
from docx import Document

# Open the DOCX file as a zip archive
with zipfile.ZipFile('example.docx') as docx_zip:

    # Find the XML file containing the document content
    document_xml_path = 'word/document.xml'
    document_xml = docx_zip.read(document_xml_path)

    # Parse the XML using lxml
    xml_tree = etree.fromstring(document_xml)

    # Define the namespace map for the XPath queries
    ns_map = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    # Find all <w:t> elements in the document
    text_elements = xml_tree.xpath('//w:t', namespaces=ns_map)

    # Print the text content of each <w:t> element
    for text_element in text_elements:
        print(text_element.text)

'''
